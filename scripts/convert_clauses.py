# -*- coding: utf-8 -*-
"""
批量将docx/pdf转换为md格式的保险条款
"""
import os
import sys

# 设置UTF-8输出
sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

def extract_docx(docx_path):
    """从docx提取文本"""
    doc = DocxDocument(docx_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 判断标题级别
            style_name = para.style.name if para.style else 'Normal'
            if 'Heading 1' in style_name or style_name == 'Title':
                paragraphs.append(f'# {text}')
            elif 'Heading 2' in style_name:
                paragraphs.append(f'## {text}')
            elif 'Heading 3' in style_name:
                paragraphs.append(f'### {text}')
            else:
                paragraphs.append(text)
    
    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                paragraphs.append('| ' + ' | '.join(cells) + ' |')
        paragraphs.append('')
    
    return '\n\n'.join(paragraphs)


def extract_pdf(pdf_path):
    """从pdf提取文本"""
    text_parts = []
    
    # 优先使用pdfplumber（更好地保留布局）
    if HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_parts.append(f'=== 第{page_num}页 ===\n{text}')
            
            if text_parts:
                return '\n\n'.join(text_parts)
        except Exception as e:
            print(f"pdfplumber failed: {e}")
    
    # 降级到pypdf
    if HAS_PYPDF:
        reader = pypdf.PdfReader(pdf_path)
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text:
                text_parts.append(f'=== 第{page_num}页 ===\n{text}')
        
        return '\n\n'.join(text_parts)
    
    return "无法读取PDF文件，请安装pdfplumber或pypdf"


def convert_to_markdown(source_path, output_path):
    """转换文件为markdown"""
    ext = os.path.splitext(source_path)[1].lower()
    
    if ext == '.docx':
        if not HAS_DOCX:
            print(f"需要安装python-docx: pip install python-docx")
            return False
        content = extract_docx(source_path)
    elif ext == '.pdf':
        content = extract_pdf(source_path)
    else:
        print(f"不支持的格式: {ext}")
        return False
    
    # 写入md文件
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    print(f"✓ 已转换: {os.path.basename(source_path)} -> {output_path}")
    return True


def main():
    """主函数"""
    downloads = r"C:\Users\庄赫\Downloads"
    temp_dir = r"C:\Users\庄赫\WorkBuddy\20260416100718\temp_clauses"
    os.makedirs(temp_dir, exist_ok=True)
    
    # 定义要转换的文件
    files = [
        ("具身智能机器人第三者责任保险条款.docx", "责任险"),
        ("机器人产品责任保险条款.docx", "责任险"),
        ("云计算服务责任保险条款.pdf", "责任险"),
        ("紫金保险附加个人意外伤害住院日额津贴保险条款.pdf", "意外险"),
        ("紫金保险附加个人疾病全残身故保险条款.pdf", "意外险"),
        ("紫金财产保险股份有限公司低速无人驾驶设备保险（2026版）条款.pdf", "特殊风险"),
        ("低速无人驾驶设备保险附加智能设备系统保险条款.pdf", "特殊风险"),
        ("机器人综合责任保险（2026版）条款.docx", "责任险"),
    ]
    
    results = []
    
    for filename, category in files:
        source = os.path.join(downloads, filename)
        output = os.path.join(temp_dir, filename.replace('.docx', '.md').replace('.pdf', '.md'))
        
        if not os.path.exists(source):
            print(f"✗ 文件不存在: {source}")
            results.append((filename, category, False, "文件不存在"))
            continue
        
        success = convert_to_markdown(source, output)
        if success:
            # 获取文件大小
            size = os.path.getsize(output)
            results.append((filename, category, True, output, size))
        else:
            results.append((filename, category, False, "转换失败"))
    
    # 输出结果
    print("\n" + "="*60)
    print("转换结果汇总")
    print("="*60)
    for r in results:
        if r[2]:  # 成功
            print(f"✓ {r[0]} ({r[1]}) - {r[4]/1024:.1f}KB")
        else:
            print(f"✗ {r[0]} ({r[1]}) - {r[3]}")
    
    return results


if __name__ == "__main__":
    main()
