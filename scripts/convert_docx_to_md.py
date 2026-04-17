# -*- coding: utf-8 -*-
"""
convert_docx_to_md.py
将 .docx 文件转换为 .md 文件，输出到同目录
用法: python convert_docx_to_md.py <file.docx>
"""
import sys
import os
import re

def docx_to_markdown(docx_path):
    """使用 python-docx 读取 .docx 并转换为 markdown"""
    from docx import Document

    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        style = para.style.name.lower()

        if "heading" in style or "title" in style or "标题" in style:
            # 提取标题级别
            level = sum(1 for c in style if c in "123456") or 1
            level = min(level, 6)
            lines.append(f"{'#' * level} {text}")
        elif "list" in style or "bullet" in style or "编号" in style or "ul" in style:
            lines.append(f"- {text}")
        elif "table" in style:
            lines.append("")
        else:
            # 正常段落，处理加粗、斜体
            text = re.sub(r'\*\*(.+?)\*\*', r'**\1**', text)
            text = re.sub(r'\*(.+?)\*', r'*\1*', text)
            lines.append(text)

    # 处理表格
    for table in doc.tables:
        lines.append("")
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)

        if not rows_data:
            continue

        # 计算每列最大宽度
        col_widths = [max(len(c) for c in col) for col in zip(*rows_data)]
        for i, row in enumerate(rows_data):
            formatted = [c.ljust(col_widths[j]) for j, c in enumerate(row)]
            lines.append("| " + " | ".join(formatted) + " |")
            if i == 0:
                lines.append("| " + " | ".join("-" * w for w in col_widths) + " |")

    return "\n".join(lines)

def main():
    if len(sys.argv) < 2:
        print("用法: python convert_docx_to_md.py <file.docx>")
        sys.exit(1)

    docx_path = sys.argv[1]
    if not os.path.exists(docx_path):
        print(f"文件不存在: {docx_path}")
        sys.exit(1)

    md_path = docx_path.replace('.docx', '.md')

    try:
        md_content = docx_to_markdown(docx_path)
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        print(f"OK: {os.path.relpath(md_path)}")
    except Exception as e:
        print(f"FAIL: {docx_path} -> {e}")
        raise

if __name__ == "__main__":
    main()
