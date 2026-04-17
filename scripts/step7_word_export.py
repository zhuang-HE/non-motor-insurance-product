#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
步骤7：输出保险条款和费率表（Word格式）
非车险产品开发工作流 - 第七步
"""

import json
import os
from typing import Dict, List, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

class WordDocumentExporter:
    """Word文档导出器"""
    
    def __init__(self):
        self.doc = None
    
    def _set_chinese_font(self, run, font_name='宋体', font_size=10.5, bold=False):
        """设置中文字体"""
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.font.bold = bold
    
    def create_clause_document(self, 
                              clause_data: Dict,
                              output_path: str) -> str:
        """创建保险条款Word文档"""
        
        doc = Document()
        
        # 设置文档默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(10.5)
        
        # 标题
        title = doc.add_heading('', level=0)
        run = title.add_run(clause_data.get('产品名称', '保险产品') + '条款')
        self._set_chinese_font(run, font_name='黑体', font_size=22, bold=True)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 条款编号
        code_para = doc.add_paragraph()
        run = code_para.add_run(f"条款编号: {clause_data.get('条款编号', '')}")
        self._set_chinese_font(run, font_size=10)
        code_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # 空行
        
        # 各章节内容
        sections = clause_data.get('章节内容', {})
        
        chapter_num = 1
        for section_name, content in sections.items():
            if not content or section_name == '短期费率表':
                continue
            
            # 章节标题
            heading = doc.add_heading('', level=1)
            run = heading.add_run(f"{self._number_to_chinese(chapter_num)}、{section_name}")
            self._set_chinese_font(run, font_name='黑体', font_size=14, bold=True)
            
            # 章节内容
            self._add_formatted_content(doc, content)
            
            chapter_num += 1
            doc.add_paragraph()  # 章节间距
        
        # 短期费率表（作为附录）
        short_term = sections.get('短期费率表', '')
        if short_term:
            heading = doc.add_heading('', level=1)
            run = heading.add_run("附录：短期费率表")
            self._set_chinese_font(run, font_name='黑体', font_size=14, bold=True)
            self._add_formatted_content(doc, short_term)
        
        # 保存文档
        doc.save(output_path)
        return output_path
    
    def create_pricing_document(self,
                               pricing_data: Dict,
                               output_path: str) -> str:
        """创建定价报告Word文档"""
        
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(10.5)
        
        # 标题
        title = doc.add_heading('', level=0)
        run = title.add_run(pricing_data.get('产品名称', '保险产品') + '定价报告')
        self._set_chinese_font(run, font_name='黑体', font_size=22, bold=True)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # 一、定价方法
        heading = doc.add_heading('', level=1)
        run = heading.add_run("一、定价方法")
        self._set_chinese_font(run, font_name='黑体', font_size=14, bold=True)
        
        method = pricing_data.get('定价方法', '')
        p = doc.add_paragraph()
        run = p.add_run(f"采用方法: {method}")
        self._set_chinese_font(run)
        
        doc.add_paragraph()
        
        # 二、基础假设
        heading = doc.add_heading('', level=1)
        run = heading.add_run("二、基础假设")
        self._set_chinese_font(run, font_name='黑体', font_size=14, bold=True)
        
        assumptions = pricing_data.get('基础假设', {})
        table = doc.add_table(rows=len(assumptions)+1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # 表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '项目'
        hdr_cells[1].text = '假设值'
        
        # 数据行
        for i, (key, value) in enumerate(assumptions.items(), 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = key
            row_cells[1].text = value
        
        doc.add_paragraph()
        
        # 三、风险因子体系
        heading = doc.add_heading('', level=1)
        run = heading.add_run("三、风险因子体系")
        self._set_chinese_font(run, font_name='黑体', font_size=14, bold=True)
        
        factors = pricing_data.get('风险因子体系', [])
        for factor in factors:
            # 因子名称
            sub_heading = doc.add_heading('', level=2)
            factor_name = factor.get('因子名称', '')
            weight = factor.get('权重', '')
            run = sub_heading.add_run(f"{factor_name}（权重{weight}）")
            self._set_chinese_font(run, font_name='黑体', font_size=12, bold=True)
            
            # 等级系数表
            levels = factor.get('等级系数', {})
            if levels:
                table = doc.add_table(rows=len(levels)+1, cols=3)
                table.style = 'Light Grid Accent 1'
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = '等级'
                hdr_cells[1].text = '系数'
                hdr_cells[2].text = '说明'
                
                for i, (level, coeff) in enumerate(levels.items(), 1):
                    row_cells = table.rows[i].cells
                    row_cells[0].text = level
                    row_cells[1].text = str(coeff)
                    row_cells[2].text = ''
            
            # 说明
            desc = factor.get('说明', '')
            if desc:
                p = doc.add_paragraph()
                run = p.add_run(f"说明: {desc}")
                self._set_chinese_font(run, font_size=9)
            
            doc.add_paragraph()
        
        # 四、定价结果
        result = pricing_data.get('定价结果', {})
        if result:
            heading = doc.add_heading('', level=1)
            run = heading.add_run("四、定价结果")
            self._set_chinese_font(run, font_name='黑体', font_size=14, bold=True)
            
            # 费率水平
            sub_heading = doc.add_heading('', level=2)
            run = sub_heading.add_run("4.1 费率水平")
            self._set_chinese_font(run, font_name='黑体', font_size=12, bold=True)
            
            rate_info = {
                '基准费率': result.get('基准费率', ''),
                '最终费率': result.get('最终费率', ''),
                '因子调整系数': result.get('因子调整系数', '')
            }
            
            table = doc.add_table(rows=len(rate_info)+1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '项目'
            hdr_cells[1].text = '数值'
            
            for i, (key, value) in enumerate(rate_info.items(), 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = key
                row_cells[1].text = value
            
            doc.add_paragraph()
            
            # 费率构成
            sub_heading = doc.add_heading('', level=2)
            run = sub_heading.add_run("4.2 费率构成")
            self._set_chinese_font(run, font_name='黑体', font_size=12, bold=True)
            
            composition = result.get('费率构成', {})
            table = doc.add_table(rows=len(composition)+1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '项目'
            hdr_cells[1].text = '比例'
            
            for i, (key, value) in enumerate(composition.items(), 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = key
                row_cells[1].text = value
            
            doc.add_paragraph()
            
            # 应用的调整因子
            sub_heading = doc.add_heading('', level=2)
            run = sub_heading.add_run("4.3 应用的调整因子")
            self._set_chinese_font(run, font_name='黑体', font_size=12, bold=True)
            
            applied = result.get('调整因子', {})
            table = doc.add_table(rows=len(applied)+1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '因子'
            hdr_cells[1].text = '系数'
            
            for i, (key, value) in enumerate(applied.items(), 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = key
                row_cells[1].text = str(value)
        
        # 保存文档
        doc.save(output_path)
        return output_path
    
    def create_complete_package(self,
                               product_name: str,
                               clause_file: str,
                               pricing_file: str,
                               output_dir: str) -> List[str]:
        """创建完整的产品文档包"""
        
        generated_files = []
        
        # 读取条款数据
        if os.path.exists(clause_file):
            with open(clause_file, 'r', encoding='utf-8') as f:
                clause_data = json.load(f)
            
            # 生成条款Word
            clause_word = os.path.join(output_dir, f"{product_name}条款.docx")
            self.create_clause_document(clause_data, clause_word)
            generated_files.append(clause_word)
        
        # 读取定价数据
        if os.path.exists(pricing_file):
            with open(pricing_file, 'r', encoding='utf-8') as f:
                pricing_data = json.load(f)
            
            # 生成定价报告Word
            pricing_word = os.path.join(output_dir, f"{product_name}定价报告.docx")
            self.create_pricing_document(pricing_data, pricing_word)
            generated_files.append(pricing_word)
        
        return generated_files
    
    def _add_formatted_content(self, doc: Document, content: str):
        """添加格式化内容"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 判断是否为条款项（以"第X条"或"（X）"开头）
            if line.startswith('第') and '条' in line[:10]:
                p = doc.add_paragraph()
                run = p.add_run(line)
                self._set_chinese_font(run, bold=True)
            elif line.startswith('（') and '）' in line[:5]:
                p = doc.add_paragraph()
                run = p.add_run(line)
                self._set_chinese_font(run)
                p.paragraph_format.left_indent = Inches(0.3)
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
                self._set_chinese_font(run)
    
    def _number_to_chinese(self, num: int) -> str:
        """数字转中文"""
        chinese_nums = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        if num <= 10:
            return chinese_nums[num]
        elif num < 20:
            return '十' + (chinese_nums[num-10] if num > 10 else '')
        elif num < 100:
            tens = num // 10
            ones = num % 10
            return chinese_nums[tens] + '十' + (chinese_nums[ones] if ones > 0 else '')
        return str(num)


def main():
    """主函数"""
    exporter = WordDocumentExporter()
    
    print("=" * 60)
    print("步骤7：输出保险条款和费率表（Word格式）")
    print("=" * 60)
    
    # 输入产品名称
    product_name = input("\n请输入产品名称: ")
    output_dir = input("请输入输出目录(默认: ./output): ") or "./output"
    
    # 创建输出目录
    os.makedirs(output_dir, exist_ok=True)
    
    # 检查输入文件
    clause_file = "step5_policy_clause.json"
    pricing_file = "step6_pricing_model.json"
    
    # 生成文档包
    generated = exporter.create_complete_package(
        product_name=product_name,
        clause_file=clause_file,
        pricing_file=pricing_file,
        output_dir=output_dir
    )
    
    if generated:
        print(f"\n✓ 文档生成成功！")
        print("生成的文件:")
        for f in generated:
            print(f"  - {f}")
    else:
        print("\n⚠ 未找到输入数据文件，请先完成步骤5和步骤6")
    
    return generated


if __name__ == "__main__":
    main()
