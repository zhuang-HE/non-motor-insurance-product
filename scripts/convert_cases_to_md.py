"""
convert_cases_to_md.py
将 cases/ 下所有子文件夹中的 .docx / .doc / .pdf 文件转换为同名 .md 文件。
同时在 cases/ 根目录生成 search_index.md，列出所有 md 文件路径，方便 AI 检索。

用法：
    python scripts/convert_cases_to_md.py
"""

import os
import sys
import re
import traceback
from pathlib import Path

# Windows PowerShell 默认 GBK，强制 stdout/stderr 用 UTF-8
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

# ── 路径 ──────────────────────────────────────────────────────────────
SKILL_ROOT = Path(__file__).resolve().parent.parent
CASES_DIR  = SKILL_ROOT / "cases"
INDEX_FILE = CASES_DIR / "search_index.md"

# ── docx 转换 ─────────────────────────────────────────────────────────
def docx_to_md(src: Path) -> str:
    """读取 .docx/.doc，返回 Markdown 字符串"""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(src))
    lines = []

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text  = para.text.strip()
        if not text:
            lines.append("")
            continue

        if "Heading 1" in style:
            lines.append(f"# {text}")
        elif "Heading 2" in style:
            lines.append(f"## {text}")
        elif "Heading 3" in style:
            lines.append(f"### {text}")
        elif "Heading 4" in style:
            lines.append(f"#### {text}")
        elif "List" in style:
            lines.append(f"- {text}")
        else:
            lines.append(text)

    # 表格
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.replace("\n", " ").strip() for c in row.cells]
            rows.append(cells)
        if not rows:
            continue
        header = rows[0]
        lines.append("")
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        for row in rows[1:]:
            # 合并单元格时可能列数不一致，补齐
            while len(row) < len(header):
                row.append("")
            lines.append("| " + " | ".join(row[:len(header)]) + " |")
        lines.append("")

    return "\n".join(lines)


# ── pdf 转换 ──────────────────────────────────────────────────────────
def pdf_to_md(src: Path) -> str:
    """读取 PDF，返回 Markdown 字符串（逐页提取文本）"""
    import pdfplumber

    pages_text = []
    with pdfplumber.open(str(src)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            # 提取表格
            tables = page.extract_tables() or []
            table_md_blocks = []
            for tbl in tables:
                if not tbl:
                    continue
                header = [str(c or "").replace("\n", " ") for c in tbl[0]]
                block  = ["| " + " | ".join(header) + " |",
                          "| " + " | ".join(["---"] * len(header)) + " |"]
                for row in tbl[1:]:
                    row = [str(c or "").replace("\n", " ") for c in row]
                    block.append("| " + " | ".join(row) + " |")
                table_md_blocks.append("\n".join(block))

            page_content = f"<!-- 第{i}页 -->\n{text}"
            if table_md_blocks:
                page_content += "\n\n" + "\n\n".join(table_md_blocks)
            pages_text.append(page_content)

    return "\n\n---\n\n".join(pages_text)


# ── 单文件转换 ────────────────────────────────────────────────────────
def convert_file(src: Path) -> Path | None:
    """将 src 转换为同目录下同名 .md 文件，返回目标路径；失败返回 None"""
    ext = src.suffix.lower()
    dst = src.with_suffix(".md")

    # 已存在则跳过（可用 --force 参数覆盖）
    force = "--force" in sys.argv
    if dst.exists() and not force:
        return dst  # already done

    try:
        if ext in (".docx", ".doc"):
            content = docx_to_md(src)
        elif ext == ".pdf":
            content = pdf_to_md(src)
        else:
            return None

        # 文件头
        header = (
            f"---\n"
            f"source_file: {src.name}\n"
            f"product: {src.parent.name}\n"
            f"---\n\n"
            f"# {src.stem}\n\n"
        )
        dst.write_text(header + content, encoding="utf-8")
        return dst

    except Exception as e:
        print(f"  [ERROR] {src.relative_to(CASES_DIR)}: {e}")
        # traceback.print_exc()
        return None


# ── 主流程 ────────────────────────────────────────────────────────────
def main():
    print(f"扫描目录: {CASES_DIR}")
    all_md_files: list[tuple[str, list[Path]]] = []   # (folder_name, [md_paths])

    total_files   = 0
    success_count = 0
    skip_count    = 0
    error_count   = 0

    for item in sorted(CASES_DIR.iterdir()):
        if not item.is_dir():
            continue

        folder_mds = []
        for src in sorted(item.iterdir()):
            if src.suffix.lower() not in (".docx", ".doc", ".pdf"):
                continue
            total_files += 1
            force = "--force" in sys.argv
            dst = src.with_suffix(".md")
            if dst.exists() and not force:
                print(f"  [SKIP]  {src.relative_to(CASES_DIR)}")
                skip_count += 1
                folder_mds.append(dst)
                continue

            print(f"  [CONV]  {src.relative_to(CASES_DIR)} ...", end=" ", flush=True)
            result = convert_file(src)
            if result:
                print("✓")
                success_count += 1
                folder_mds.append(result)
            else:
                print("✗")
                error_count += 1

        if folder_mds:
            all_md_files.append((item.name, folder_mds))

    # 同级 md 文件（cases_index.md 等）不转换，直接收录
    # 写搜索索引
    print(f"\n生成索引: {INDEX_FILE}")
    lines = [
        "# 非车险案例库检索索引\n",
        "> 本文件由脚本自动生成。列出 cases/ 下所有 Markdown 文件路径，供 AI Agent 快速定位。\n",
        f"**总产品数：{len(all_md_files)}  总文档数：{total_files}**\n",
        "---\n",
    ]
    for folder, mds in all_md_files:
        lines.append(f"\n## {folder}\n")
        for md in mds:
            rel = md.relative_to(CASES_DIR)
            # 文件名去掉后缀作为简称
            label = md.stem
            lines.append(f"- [{label}]({rel.as_posix()})")

    INDEX_FILE.write_text("\n".join(lines), encoding="utf-8")

    print(f"\n完成：转换 {success_count}，跳过 {skip_count}，失败 {error_count}，共 {total_files} 个文件")
    print(f"索引已写入: {INDEX_FILE}")


if __name__ == "__main__":
    main()
